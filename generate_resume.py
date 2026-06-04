"""Generate a one-page resume (.docx) for Sai Abhinav.

Run:
    venv\\Scripts\\python.exe generate_resume.py
Output:
    Sai_Abhinav_Resume.docx
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "Sai_Abhinav_Resume.docx"
FONT = "Calibri"
ACCENT = RGBColor(0x1F, 0x36, 0x5F)  # dark navy


def set_run(run, *, size=10, bold=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def tight(paragraph, *, before=0, after=0, line=1.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F365F")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def section_header(doc, text):
    p = doc.add_paragraph()
    tight(p, before=8, after=4, line=1.0)
    add_bottom_border(p)
    r = p.add_run(text.upper())
    set_run(r, size=12, bold=True, color=ACCENT)
    r.font.name = FONT
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    tight(p, before=1, after=1, line=1.15)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    # Clear default run and re-add to control font
    for r in p.runs:
        r.text = ""
    r = p.add_run(text)
    set_run(r, size=11)
    return p


def add_line(doc, runs, *, before=0, after=0, align=None):
    p = doc.add_paragraph()
    tight(p, before=before, after=after, line=1.15)
    if align is not None:
        p.alignment = align
    for spec in runs:
        text = spec["text"]
        r = p.add_run(text)
        set_run(
            r,
            size=spec.get("size", 11),
            bold=spec.get("bold", False),
            color=spec.get("color"),
        )
    return p


def build():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    # ── Header ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(name, before=0, after=2)
    r = name.add_run("SAI ABHINAV")
    set_run(r, size=26, bold=True, color=ACCENT)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(contact, before=0, after=6)
    r = contact.add_run(
        "+91-9958434717  •  saiabhinav190404@gmail.com  •  "
        "github.com/SaixAbhinav  •  Noida, UP"
    )
    set_run(r, size=10.5)

    # ── Summary ─────────────────────────────────────────────────────────────
    section_header(doc, "Summary")
    p = doc.add_paragraph()
    tight(p, before=2, after=0, line=1.2)
    r = p.add_run(
        "Applied AI builder shipping automation-driven solutions across ML, RL, "
        "and LLM workflows. Experienced in real-world systems including AI-powered "
        "traffic optimization, detection models, and local-first LLM tooling. "
        "Strong interest in workflow automation, prompt engineering, and deploying "
        "AI to solve operational problems."
    )
    set_run(r, size=11)

    # ── Projects ────────────────────────────────────────────────────────────
    section_header(doc, "Projects")

    add_line(
        doc,
        [
            {"text": "AI Workflow Copilot ", "bold": True, "size": 11.5},
            {"text": "— Local-First Desktop Assistant", "size": 11.5},
        ],
        before=4,
        after=0,
    )
    add_bullet(
        doc,
        "PyQt5 desktop app that converts documents (PDF/TXT) and Gmail threads into "
        "summaries, action items, insights, and side-by-side comparisons via a local "
        "Ollama LLM — fully offline, no data leaves the machine.",
    )
    add_bullet(
        doc,
        "Designed a modular pipeline (input handlers → chunker → prompt manager → "
        "LLM router → JSON parser → result merger) supporting four distinct workflows "
        "behind a single orchestrator.",
    )
    add_bullet(
        doc,
        "Integrated Gmail and Google Calendar APIs with multi-account OAuth, a "
        "task-review dialog with editable deadlines and priorities, SQLite-backed "
        "run history, and TXT/CSV export.",
    )

    add_line(
        doc,
        [
            {"text": "SmartSignal ", "bold": True, "size": 11.5},
            {"text": "— AI Traffic Light Automation System", "size": 11.5},
        ],
        before=6,
        after=0,
    )
    add_bullet(
        doc,
        "Built a PPO-based traffic signal controller in SUMO simulation that reacts "
        "to live conditions, reducing vehicle wait time by ~28% and lifting throughput "
        "by 22%.",
    )
    add_bullet(
        doc,
        "Integrated the TomTom API for realistic traffic flow modelling and designed "
        "the system to scale across multiple junctions simultaneously.",
    )

    add_line(
        doc,
        [
            {"text": "Instagram Fake Profile Detection System", "bold": True, "size": 11.5},
        ],
        before=6,
        after=0,
    )
    add_bullet(
        doc,
        "Built an ensemble model combining classical ML with CNN-based image "
        "classification on 3,000+ accounts, reaching 92% accuracy, 90% precision, "
        "and 88% recall.",
    )
    add_bullet(
        doc,
        "Tuned thresholds and feature set to reduce false positives by 15%.",
    )

    # ── Experience ──────────────────────────────────────────────────────────
    section_header(doc, "Experience")
    add_line(
        doc,
        [
            {"text": "IBM Skills Build ", "bold": True, "size": 11.5},
            {"text": "— Data Analyst Intern", "size": 11.5},
            {"text": "     Jun 2024 – Aug 2024", "size": 10.5},
        ],
        before=2,
        after=0,
    )
    add_bullet(
        doc,
        "Built and optimized an image-based skin cancer detection model, reaching "
        "94% accuracy through preprocessing, feature engineering, and model tuning.",
    )
    add_bullet(
        doc,
        "Collaborated in a 6-member team; presented findings and recommendations "
        "to stakeholders.",
    )

    # ── Education ───────────────────────────────────────────────────────────
    section_header(doc, "Education")
    add_line(
        doc,
        [
            {"text": "Vivekananda Institute of Professional Studies, Delhi", "bold": True, "size": 11},
            {"text": "  — MCA, 2025–2027  ·  CGPA 8.6", "size": 11},
        ],
        before=2,
        after=0,
    )
    add_line(
        doc,
        [
            {"text": "Vivekananda Institute of Professional Studies, Delhi", "bold": True, "size": 11},
            {"text": "  — BCA, 2022–2025  ·  CGPA 7.6", "size": 11},
        ],
        before=2,
        after=0,
    )

    # ── Skills ──────────────────────────────────────────────────────────────
    section_header(doc, "Skills")

    def skill_line(label, value):
        p = doc.add_paragraph()
        tight(p, before=2, after=0, line=1.15)
        r = p.add_run(f"{label}:  ")
        set_run(r, size=11, bold=True)
        r = p.add_run(value)
        set_run(r, size=11)

    skill_line("Programming", "Python, SQL")
    skill_line(
        "AI / ML",
        "Supervised & Unsupervised Learning, Feature Engineering, Model Optimization, Reinforcement Learning (PPO)",
    )
    skill_line("Frameworks", "TensorFlow, Keras, NumPy, Pandas, Flask, PyQt5")
    skill_line(
        "AI Tooling",
        "OpenAI API, Ollama (local LLMs), Prompt Engineering, Workflow Automation",
    )
    skill_line("Data & Viz", "Matplotlib, Seaborn, SQLite, Excel / Google Sheets")

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
