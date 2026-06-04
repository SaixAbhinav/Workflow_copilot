from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

emails = [
    {
        "title": "Email 1 — Project status update (Summary + Tasks)",
        "headers": [
            ("From", "priya.menon@northwind.com"),
            ("To", "team-platform@northwind.com"),
            ("Subject", "Q2 Migration — week 3 status & next steps"),
        ],
        "body": [
            "Hi team,",
            "Quick update on the Postgres to Aurora migration. We finished the schema diff this week and the read-replica is in place in staging. Latency on the heavy analytics queries dropped ~38% which is better than the 25% we projected.",
            "A few items still open:",
            "  - Rahul to finish the connection-pool config by Friday May 22",
            "  - Sam to run the failover drill next Tuesday at 10am",
            "  - I'll send the rollback runbook for review by EOW",
            "Risks: the legacy reporting job still uses a hardcoded DSN. We need to patch that before cutover on June 3.",
            "Please RSVP to the cutover dry-run invite by Monday.",
            "Thanks,",
            "Priya",
        ],
    },
    {
        "title": "Email 2 — Meeting invite (Calendar push)",
        "headers": [
            ("From", "alex.huang@northwind.com"),
            ("To", "saiabhinav190404@gmail.com"),
            ("Subject", "Design review — Workflow Copilot UX, Thurs May 21, 3:00–4:00 PM IST"),
        ],
        "body": [
            "Hi Sai,",
            "Locking in our design review for Thursday May 21, 3:00–4:00 PM IST. We'll cover the new Compare panel layout and the keyboard-shortcut flow for the Tasks view.",
            "Agenda:",
            "  1. Walk through Figma v4 (10 min)",
            "  2. Open questions on the diff highlighting (20 min)",
            "  3. Accessibility audit findings (15 min)",
            "  4. Action items + owners (15 min)",
            "Please review the Figma link before the call so we can spend the time on decisions, not catch-up.",
            "Zoom: https://zoom.us/j/9988776655",
            "— Alex",
        ],
    },
    {
        "title": "Email 3 — Customer feedback (Insights)",
        "headers": [
            ("From", "dana.fields@acmeretail.com"),
            ("To", "support@northwind.com"),
            ("Subject", "Honest feedback after 30 days using Copilot"),
        ],
        "body": [
            "Hi team,",
            "Wanted to share what my team's experience has been after a month. Overall — the summarization quality is the standout feature. Three of my analysts independently said it's saving them 4–6 hours a week on report digestion.",
            "What's working well:",
            "  - Summary accuracy on 20+ page PDFs is very strong",
            "  - Local-only processing is a hard requirement for us (legal), and it's been smooth",
            "  - The Tasks extraction caught a deadline two people had missed in a contract",
            "What's not:",
            "  - Compare view feels cluttered when documents are >50 pages — would love collapsible sections",
            "  - Gmail import sometimes misses threads older than 90 days",
            "  - Insights occasionally repeats the same point twice in different wording",
            "Happy to jump on a call if useful. We're renewing for sure.",
            "Dana",
            "Director of Operations, Acme Retail",
        ],
    },
    {
        "title": "Email 4 — Contract revision (Compare)",
        "headers": [
            ("From", "legal@brightpath-partners.com"),
            ("To", "saiabhinav190404@gmail.com"),
            ("Subject", "Revised MSA — v3 attached, please review by May 20"),
        ],
        "body": [
            "Hi Sai,",
            "Attached is v3 of the Master Services Agreement. Changes from v2:",
            "  1. Payment terms moved from Net-45 to Net-30",
            "  2. Liability cap raised from $50,000 to $150,000",
            "  3. Added a data-processing addendum (Schedule C) — required for our EU clients",
            "  4. Termination-for-convenience window changed from 60 days to 90 days",
            "  5. Removed the auto-renewal clause entirely",
            "Item 4 is the one most likely to need discussion on your side. Everything else aligns with what we agreed to verbally on the May 8 call.",
            "Please send signed copy or redlines by Tuesday May 20.",
            "Best,",
            "Marcus Hale",
            "General Counsel, BrightPath Partners",
        ],
    },
    {
        "title": "Email 5 — Daily ops digest (Tasks + Insights stress test)",
        "headers": [
            ("From", "ops-daily@northwind.com"),
            ("To", "leadership@northwind.com"),
            ("Subject", "Daily ops digest — May 15"),
        ],
        "body": [
            "Morning all,",
            "Incidents (last 24h):",
            "  - SEV-3 at 02:14 UTC: queue worker backlog hit 12k jobs; auto-scaled, cleared in 27 min. Postmortem owner: Jen.",
            "  - Auth latency spike at 09:40 UTC — root cause was a deploy from the staging branch hitting prod by mistake. Reverted. Need to discuss branch protections this week.",
            "Hires: Two offers out (Senior SRE, Product Designer). Both expected to respond by Friday.",
            "Decisions needed today:",
            "  1. Approve the $18k spend on the observability vendor renewal — Karthik blocked on this.",
            "  2. Pick a date for the all-hands: June 4 or June 11. Marketing prefers June 11.",
            "  3. Sign off on the new on-call rotation (draft in Notion).",
            "Wins: Mobile crash rate down 22% week-over-week after the v4.2 patch. NPS jumped 6 points in the latest cohort survey.",
            "Reply-all if you need to escalate anything before standup at 11.",
            "— Ops desk",
        ],
    },
]

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_heading("Demo Emails — Workflow Copilot", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

intro = doc.add_paragraph()
intro.add_run(
    "Five sample emails for demoing Summary, Tasks, Insights, Compare, and Calendar push features."
).italic = True
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

for i, email in enumerate(emails):
    doc.add_heading(email["title"], level=1)

    for label, value in email["headers"]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    doc.add_paragraph()

    for line in email["body"]:
        doc.add_paragraph(line)

    if i < len(emails) - 1:
        doc.add_paragraph("—" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

out_path = "demo_emails.docx"
doc.save(out_path)
print(f"Saved {out_path}")
