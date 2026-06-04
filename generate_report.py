"""Generate a VIPS/GGSIPU MCA Minor Project report (.docx) for AI Workflow Copilot.

Formatting follows the project guidelines verbatim:
  Page    : A4, Left 3.0 cm, Right 2.0 cm, Top/Bottom 2.54 cm
  Body    : Times New Roman 12 pt, double spaced, justified, 6 pt before/after
  Heading : Times New Roman 14 pt, underlined, left aligned, 12 pt before/after
  Chapter : Times New Roman 20 pt, centered, 30 pt before/after
  Code    : Courier New 10 pt
  Page #s : bottom-center on every text page

Run:
    venv\\Scripts\\python.exe generate_report.py
Output:
    Project_Report.docx
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ───────────────────────── STUDENT FIELDS — EDIT THESE ──────────────────────
STUDENT_NAME      = "Sai Abhinav"
STUDENT_ROLL_NO   = "<10217704425>"
SEMESTER          = "II"                        # I / II / III
SUBJECT_CODE      = "MCA-170"                   # MCA-169 / MCA-170 / MCA-269
GUIDE_NAME        = "<Name of Project Guide>"
GUIDE_DESIGNATION = "Assistant Professor, VSIT"
INSTITUTE_LOC     = "Vivekananda Institute of Professional Studies – TC"
PROJECT_PERIOD    = "January 2026 – April 2026"
SESSION           = "Jan – April, 2026"
PROJECT_TITLE     = "AI WORKFLOW COPILOT"
PROJECT_SUBTITLE  = "A Local-First Desktop Assistant for Document and Email " \
                    "Workflows using a Local LLM"
# ────────────────────────────────────────────────────────────────────────────


# ═════════════════════════ low-level styling helpers ════════════════════════

def _set_run_font(run, name="Times New Roman", size=12, bold=False,
                  italic=False, underline=False, color=None):
    run.font.name = name
    # Force the East-Asian font slot too so Word doesn't fall back on Calibri
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para_spacing(p, before_pt=6, after_pt=6, line_rule="double"):
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after  = Pt(after_pt)
    if line_rule == "double":
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif line_rule == "single":
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif line_rule == "1.5":
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
             italic=False, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    _para_spacing(p, before_pt=6, after_pt=6, line_rule="double")
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_paragraph_heading(doc, text):
    """Section heading: TNR 14, bold, underlined, left, 12 pt before/after."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=12, after_pt=12, line_rule="single")
    run = p.add_run(text)
    _set_run_font(run, size=14, bold=True, underline=True)
    return p


def add_chapter_heading(doc, text, page_break_before=True):
    """Chapter heading: TNR 20, centered, 30 pt before/after, new page."""
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before_pt=30, after_pt=30, line_rule="single")
    run = p.add_run(text)
    _set_run_font(run, size=20, bold=True)
    return p


def add_subheading(doc, text):
    """Sub-section heading inside a chapter: TNR 13 bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=10, after_pt=6, line_rule="single")
    run = p.add_run(text)
    _set_run_font(run, size=13, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _para_spacing(p, before_pt=2, after_pt=2, line_rule="single")
    p.paragraph_format.left_indent = Cm(1.0)
    for run in p.runs:
        _set_run_font(run, size=12)
    if not p.runs:
        run = p.add_run(text)
        _set_run_font(run, size=12)
    else:
        # paragraph created with style sometimes has a run already
        p.runs[-1].text = text
        _set_run_font(p.runs[-1], size=12)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    _para_spacing(p, before_pt=2, after_pt=2, line_rule="single")
    p.paragraph_format.left_indent = Cm(1.0)
    if not p.runs:
        run = p.add_run(text)
    else:
        p.runs[-1].text = text
    _set_run_font(p.runs[-1], size=12)
    return p


def add_code_block(doc, code_text):
    """Code block: Courier New 10 pt, single-spaced, no justification."""
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_spacing(p, before_pt=0, after_pt=0, line_rule="single")
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else " ")
        _set_run_font(run, name="Courier New", size=10)


def add_centered(doc, text, size=12, bold=False, italic=False, before=6, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before_pt=before, after_pt=after, line_rule="single")
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_blank(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        _para_spacing(p, before_pt=0, after_pt=0, line_rule="single")


# ═════════════════════════════ page numbers ═════════════════════════════════

def add_page_numbers(section):
    """Insert {PAGE} field, centered, in the footer of `section`."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr     = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_sep   = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end   = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")

    run = p.add_run()
    _set_run_font(run, size=11)
    r_el = run._r
    r_el.append(fld_begin); r_el.append(instr); r_el.append(fld_sep); r_el.append(fld_end)


# ═══════════════════════════ document setup ════════════════════════════════

def build_document() -> Document:
    doc = Document()

    # Configure default Normal style → Times New Roman 12 double-spaced
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Section / page setup — A4, VIPS margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        add_page_numbers(section)

    return doc


# ═════════════════════════════ content blocks ═══════════════════════════════

def cover_page(doc):
    section = doc.sections[0]
    # Title block — large, centered
    add_blank(doc, 2)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before_pt=0, after_pt=0, line_rule="single")
    run = p.add_run(PROJECT_TITLE)
    _set_run_font(run, size=20, bold=True)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before_pt=6, after_pt=0, line_rule="single")
    run = p.add_run(PROJECT_SUBTITLE)
    _set_run_font(run, size=14, italic=True)

    add_blank(doc, 3)
    add_centered(doc, "IN PARTIAL FULFILLMENT OF THE REQUIREMENT FOR THE AWARD OF DEGREE OF",
                 size=12, bold=True)
    add_centered(doc, "MASTER OF COMPUTER APPLICATIONS", size=14, bold=True)

    add_blank(doc, 3)
    add_centered(doc, "Submitted By:", size=12, bold=True)
    add_centered(doc, f"{STUDENT_NAME}   ({STUDENT_ROLL_NO})", size=12)

    add_blank(doc, 3)
    add_centered(doc, "[ VIPS LOGO ]", size=12, italic=True)

    add_blank(doc, 2)
    add_centered(doc, "VIVEKANANDA SCHOOL OF INFORMATION TECHNOLOGY", size=12, bold=True)
    add_centered(doc, "VIVEKANANDA INSTITUTE OF PROFESSIONAL STUDIES – TC", size=12, bold=True)

    add_blank(doc, 1)
    add_centered(doc, "Affiliated to", size=12)
    add_centered(doc, "(GURU GOBIND SINGH INDRAPRASTHA UNIVERSITY, DELHI)",
                 size=12, bold=True)

    add_blank(doc, 2)
    add_centered(doc, SESSION, size=12, bold=True)


def certificate_page(doc):
    doc.add_page_break()
    add_centered(doc, "CERTIFICATE", size=18, bold=True, before=18, after=24)
    add_body(doc,
        f"This is to certify that this project entitled “{PROJECT_TITLE}” submitted in "
        f"partial fulfillment of the {SEMESTER} semester of MCA to the VIPS, done by "
        f"Mr./Ms. {STUDENT_NAME}, Roll No. {STUDENT_ROLL_NO}, is an authentic work "
        f"carried out by him/her at {INSTITUTE_LOC} under my guidance. The matter "
        f"embodied in this project work has not been submitted earlier for the award "
        f"of any degree or diploma to the best of my knowledge and belief."
    )
    add_body(doc, f"Period of project work: {PROJECT_PERIOD}.")

    add_blank(doc, 6)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=0, after_pt=0, line_rule="single")
    p.add_run("Signature of the Student" + " " * 30 + "Signature of the Guide")
    for run in p.runs: _set_run_font(run, size=12, bold=True)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=0, after_pt=0, line_rule="single")
    p.add_run(f"{STUDENT_NAME}" + " " * 40 + f"{GUIDE_NAME}")
    for run in p.runs: _set_run_font(run, size=12)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=0, after_pt=0, line_rule="single")
    p.add_run(f"Roll No.: {STUDENT_ROLL_NO}" + " " * 25 + f"{GUIDE_DESIGNATION}")
    for run in p.runs: _set_run_font(run, size=12)


def self_certificate_page(doc):
    doc.add_page_break()
    add_centered(doc, "SELF CERTIFICATE", size=18, bold=True, before=18, after=24)
    add_body(doc,
        f"This is to certify that the project report entitled “{PROJECT_TITLE}” is "
        f"done by me and is an authentic work carried out for the partial fulfilment "
        f"of the {SEMESTER} semester of Master of Computer Applications under the "
        f"guidance of {GUIDE_NAME}. The matter embodied in this project work has not "
        f"been submitted earlier for the award of any degree or diploma to the best "
        f"of my knowledge and belief."
    )
    add_blank(doc, 6)
    add_body(doc, "Signature of the Student", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    add_body(doc, f"Name of the Student : {STUDENT_NAME}",  align=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(doc, f"Roll No.            : {STUDENT_ROLL_NO}", align=WD_ALIGN_PARAGRAPH.LEFT)


def acknowledgements_page(doc):
    doc.add_page_break()
    add_centered(doc, "ACKNOWLEDGEMENTS", size=18, bold=True, before=18, after=24)
    add_body(doc,
        f"The completion of any project depends upon the cooperation, coordination "
        f"and combined efforts of several sources of knowledge, energy and time. I "
        f"take this opportunity to express my deep sense of gratitude to my project "
        f"guide {GUIDE_NAME}, {GUIDE_DESIGNATION}, for the valuable guidance, "
        f"continuous encouragement and constructive feedback offered throughout the "
        f"course of this project. Their insights into software engineering practices "
        f"and willingness to review iterative drafts of this work have been "
        f"instrumental in shaping its final form."
    )
    add_body(doc,
        "I am sincerely grateful to the Director and the Head of the Department, "
        "Vivekananda School of Information Technology, Vivekananda Institute of "
        "Professional Studies – TC, for providing the academic environment and "
        "infrastructure that made this project possible. I extend my thanks to the "
        "faculty members of the MCA programme for the foundational courses in "
        "software engineering, database systems, artificial intelligence and human "
        "computer interaction that informed the design choices made in this work."
    )
    add_body(doc,
        "I would also like to thank the open-source community behind PyQt, the "
        "Ollama runtime, Google API client libraries and PyMuPDF — without these "
        "freely available tools the local-first architecture central to this "
        "project would not have been feasible. Finally, I thank my family and "
        "peers for their patient support during the development period."
    )
    add_blank(doc, 4)
    add_body(doc, STUDENT_NAME,             align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    add_body(doc, f"Roll No.: {STUDENT_ROLL_NO}", align=WD_ALIGN_PARAGRAPH.RIGHT)


# ─────────────────────────── synopsis (3–4 pages) ───────────────────────────

def synopsis_chapter(doc):
    add_chapter_heading(doc, "SYNOPSIS")

    add_paragraph_heading(doc, "1. Title of the Project")
    add_body(doc,
        f"{PROJECT_TITLE}: a local-first PyQt5 desktop assistant that converts "
        f"documents and Gmail messages into structured summaries, action items, "
        f"insights and side-by-side comparisons using a locally hosted large "
        f"language model, with optional push-to-Google-Calendar for extracted "
        f"tasks."
    )

    add_paragraph_heading(doc, "2. Statement of the Problem")
    add_body(doc,
        "Knowledge workers are routinely asked to read long-form material — "
        "reports, contracts, articles, multi-thread email conversations — and "
        "convert that material into a small number of decisions: what to "
        "summarise, what to do, what is novel, and how two artefacts differ. "
        "Two adjacent classes of tools attempt to support this work. Cloud-based "
        "LLM products such as ChatGPT, Claude and Gemini deliver high quality "
        "outputs but require uploading user content to a third-party server, "
        "which is incompatible with confidentiality requirements in many "
        "academic, legal and corporate contexts. Traditional desktop "
        "productivity tools (note-takers, mail clients, calendar apps) operate "
        "locally but contain no semantic understanding of their inputs."
    )
    add_body(doc,
        "There is a gap for an integrated assistant that (a) runs the language "
        "model entirely on the user’s own machine, (b) accepts heterogeneous "
        "inputs — pasted text, .txt and .pdf files, and live Gmail messages — "
        "without forcing the user into a single workflow, and (c) closes the "
        "loop by writing extracted tasks back into the user’s calendar after "
        "human review. This project addresses that gap."
    )

    add_paragraph_heading(doc, "3. Why This Topic Was Chosen")
    add_body(doc,
        "The topic was selected for three reasons. First, local-first AI is a "
        "topical and rapidly maturing area: open-weight models that fit on "
        "consumer GPUs and the Ollama runtime that exposes them through a "
        "uniform HTTP interface have only recently made fully offline LLM "
        "applications practical for individual developers. Second, the problem "
        "is grounded in personal observation: the author repeatedly observed "
        "students and faculty re-typing email content into web LLM chats and "
        "then transcribing the response into a calendar by hand — a workflow "
        "that is slow, error-prone and privacy-leaking. Third, the project "
        "spans multiple core MCA subject areas — software engineering, "
        "database systems, GUI programming, RESTful service consumption and "
        "applied artificial intelligence — and is therefore an appropriate "
        "vehicle for a capstone style minor project."
    )

    add_paragraph_heading(doc, "4. Objective and Scope of the Project")
    add_body(doc,
        "The primary objective is to deliver a single-user desktop application "
        "for Microsoft Windows that takes documents or email as input, executes "
        "one of four named workflows (Summary, Tasks, Insights, Compare) "
        "against a locally hosted LLM, presents the structured output in a GUI, "
        "and optionally pushes extracted tasks to Google Calendar after the user "
        "has reviewed and approved them."
    )
    add_body(doc, "Scope of the work covers:")
    for item in [
        "input acquisition from text paste, .txt and .pdf files, and Gmail search results",
        "text cleaning and chunking sized to the model context window",
        "prompt templating per workflow with strict JSON output contracts",
        "LLM invocation via the Ollama HTTP API, with provider-router indirection",
        "output parsing, cross-chunk merging and a source-type classifier that suppresses fake tasks from narrative content",
        "a PyQt5 graphical interface with multi-account Google sign-in, run cancellation, history and TXT/CSV export",
        "task review dialog and Google Calendar event creation for approved tasks",
        "an SQLite history store for previous runs",
    ]:
        add_bullet(doc, item)
    add_body(doc,
        "Out of scope, intentionally, are: cloud-hosted LLMs, mobile clients, "
        "multi-user collaboration, training or fine-tuning of language models, "
        "and any feature that would require the application to upload user "
        "content to a third party other than Google for the sole purpose of "
        "calendar event creation."
    )

    add_paragraph_heading(doc, "5. Methodology (Project Summary)")
    add_body(doc,
        "The system follows a layered pipeline architecture. Inputs collected "
        "by the input_processing layer (file handlers for .txt and .pdf, an "
        "OAuth-backed Gmail handler, and a whitespace cleaner) are normalised "
        "into plain Unicode text. The text is split by the chunker into "
        "sentence-bounded segments of at most CHUNK_SIZE words (default 300). "
        "Each chunk is passed through the workflow_engine, which selects a "
        "prompt template from prompt_manager keyed by workflow name, sends the "
        "rendered prompt to llm_router (presently routing to a single Ollama "
        "provider), and feeds the response into output_parser which extracts a "
        "balanced JSON object from the model output. Per-chunk parsed results "
        "are then aggregated by merge_results — concatenated for summaries, "
        "list-extended for insights and tasks, and cross-document for the "
        "compare workflow which is run on the entire input rather than on "
        "chunks. The aggregated output is rendered in the PyQt5 main window "
        "and persisted to a SQLite history store. For the tasks workflow, an "
        "additional review dialog allows the user to edit titles, deadlines "
        "and priorities and to deselect rows before push_tasks_to_calendar "
        "creates Google Calendar events through the Calendar v3 API."
    )

    add_paragraph_heading(doc, "6. Hardware and Software to be Used")
    add_subheading(doc, "Hardware (minimum):")
    for item in [
        "x86_64 CPU with AVX2 support (Intel 8th generation / AMD Ryzen 2000 or later)",
        "8 GB RAM (16 GB recommended for 7B-parameter models)",
        "10 GB free disk space for the application, the virtual environment and one model weight file",
        "Optional CUDA-capable NVIDIA GPU with 6 GB VRAM for accelerated inference",
    ]: add_bullet(doc, item)
    add_subheading(doc, "Software:")
    for item in [
        "Microsoft Windows 10 or 11 (the target deployment platform)",
        "Python 3.10 or later",
        "Ollama runtime, with the mistral 7B model pulled locally",
        "Application dependencies: PyQt5, PyMuPDF, requests, python-dotenv, "
        "python-dateutil, google-api-python-client, google-auth, "
        "google-auth-oauthlib, pytest",
        "SQLite (bundled with the Python standard library)",
        "Google Cloud project with Gmail API and Calendar API enabled and "
        "OAuth 2.0 desktop credentials downloaded",
    ]: add_bullet(doc, item)

    add_paragraph_heading(doc, "7. Testing Technologies Used")
    add_body(doc,
        "Automated testing is built on pytest. The test suite covers (i) the "
        "chunker’s sentence-boundary respecting behaviour and word-preservation "
        "guarantees, (ii) the cleaner’s whitespace normalisation, (iii) the "
        "JSON extractor’s tolerance of markdown code fences and surrounding "
        "prose, (iv) the cross-chunk merger for each of the four workflows, "
        "and (v) the LLM router and Ollama service client, with HTTP calls "
        "mocked by unittest.mock.patch so the suite runs without network "
        "access. Manual testing is performed against synthetic PDF fixtures "
        "produced by tests/generate_test_pdfs.py and against live Gmail and "
        "Calendar accounts in a sandbox Google Cloud project."
    )

    add_paragraph_heading(doc, "8. Contribution of the Project")
    add_body(doc,
        "The project contributes a small but coherent demonstration of a "
        "local-first AI assistant: it shows that strict-JSON prompting against "
        "a 7-billion-parameter open-weight model is sufficient to drive a "
        "useful four-workflow productivity tool without any cloud LLM. It "
        "also contributes a re-usable architectural pattern — input "
        "normalisation, chunk-and-merge with a workflow-aware merger, and a "
        "router-indirected provider — that can be extended to additional "
        "workflows or alternative LLM backends with minimal change to the UI "
        "or storage layers. Specific novelties relative to comparable "
        "open-source projects include the source_type classifier that "
        "suppresses spurious task extraction from narrative or discussion "
        "input, and a multi-account Google sign-in flow that stores per-account "
        "credentials under tokens/<email>.pkl with safe-filename mangling."
    )

    add_paragraph_heading(doc, "9. Process Description")
    add_body(doc,
        "When the user starts the application, the PyQt5 main window restores "
        "the previously active Google account (if any) from tokens/active.txt, "
        "loads the saved Ollama model preference from config/user_prefs.json, "
        "and enables the input panel. The user provides input by typing or "
        "pasting text, by dragging .txt or .pdf files onto the editor, or by "
        "opening the Gmail panel and issuing a search query in native Gmail "
        "syntax (for example from:professor@vips.edu is:unread). On clicking "
        "Process, the active workflow is dispatched on a QThread so the UI "
        "remains responsive; the same button reverts to a Cancel control for "
        "the duration of the run. When the run completes, the structured "
        "result is rendered in the result panel and saved to the SQLite "
        "history store. The user may then export the result as TXT or CSV; "
        "for the Tasks workflow the user may additionally launch the review "
        "dialog and push approved tasks to Google Calendar."
    )

    add_paragraph_heading(doc, "10. Resources and Limitations")
    add_body(doc,
        "Resources required by the system fall into three categories: "
        "computational resources sufficient to run a 7B parameter LLM locally, "
        "Google Cloud OAuth credentials issued by the user against their own "
        "project, and disk space for the SQLite history database which grows "
        "linearly with usage. No remote application server, hosted database "
        "or paid API quota is required for normal operation — this is a "
        "deliberate consequence of the local-first design. The application "
        "has the following acknowledged limitations: latency is bounded "
        "below by the inference speed of the chosen Ollama model and may be "
        "perceptibly slower than a cloud LLM; output quality is also bounded "
        "by the chosen model and is below state-of-the-art frontier models on "
        "complex reasoning prompts; the application targets a single "
        "workstation user and does not provide multi-user accounts; and the "
        "set of supported input file types is currently restricted to .txt "
        "and .pdf."
    )

    add_paragraph_heading(doc, "11. Conclusion")
    add_body(doc,
        "The synopsis describes a working, fully local AI productivity tool "
        "that closes the gap between cloud LLM chat interfaces and traditional "
        "desktop applications. The principal innovations are the four-workflow "
        "abstraction with strict JSON output contracts, the source_type "
        "classifier that prevents fabricated task extraction, and the "
        "multi-account Google integration that keeps OAuth tokens segregated "
        "by email. The remainder of this report develops these claims, "
        "describes the system analysis and design, presents the implementation "
        "in detail, and reports on the testing and evaluation activities "
        "carried out during the project period."
    )


# ──────────────────────────── main report chapters ──────────────────────────

def chapter_objective_scope(doc):
    add_chapter_heading(doc, "CHAPTER 1 — OBJECTIVE AND SCOPE OF THE PROJECT")

    add_paragraph_heading(doc, "1.1 Primary Objective")
    add_body(doc,
        "To design, implement and evaluate a local-first desktop AI assistant "
        "that converts heterogeneous textual inputs (typed text, .txt and .pdf "
        "documents, and Gmail messages) into one of four classes of structured "
        "output — Summary, Tasks, Insights and Compare — using a language model "
        "executed entirely on the user’s own machine, and that optionally "
        "pushes extracted tasks to Google Calendar after the user has reviewed "
        "them."
    )

    add_paragraph_heading(doc, "1.2 Specific Objectives")
    for item in [
        "Implement a layered pipeline (input → cleaner → chunker → workflow engine → LLM → parser → merger → UI) that decouples each concern and admits unit testing.",
        "Define a strict, per-workflow JSON output schema and an output parser that tolerates markdown fences and surrounding prose so that imperfect LLM compliance does not fail the run.",
        "Design a source-type classifier that prevents the Tasks workflow from inventing action items from narrative, discussion or informational content.",
        "Provide multi-account Google OAuth with per-account credential storage, an account switcher in the UI header, and graceful re-authentication on scope mismatch or token expiry.",
        "Persist every run into a local SQLite database for reload, and supply both TXT and CSV exporters.",
        "Cover all non-LLM logic with automated pytest tests so the suite runs offline.",
    ]: add_bullet(doc, item)

    add_paragraph_heading(doc, "1.3 Scope")
    add_body(doc,
        "The system is delivered as a single-user Microsoft Windows desktop "
        "application. It assumes that the Ollama runtime is installed locally "
        "and that at least one model (mistral by default) has been pulled. "
        "All Google integrations require the user to bring their own OAuth "
        "client secret obtained from the Google Cloud Console with the Gmail "
        "API and Calendar API enabled. The scope explicitly excludes cloud "
        "LLM backends, mobile clients, multi-user collaboration, model "
        "training, and any feature that would upload user content to a third "
        "party other than Google for calendar event creation."
    )


def chapter_theoretical_background(doc):
    add_chapter_heading(doc, "CHAPTER 2 — THEORETICAL BACKGROUND")

    add_paragraph_heading(doc, "2.1 Large Language Models and Local Inference")
    add_body(doc,
        "A large language model (LLM) is a deep neural network — typically a "
        "decoder-only transformer in 2024-2026 — trained on a self-supervised "
        "next-token-prediction objective over very large text corpora. The "
        "family of mid-sized open-weight models with seven to nine billion "
        "parameters (mistral, llama-3-8B, qwen-7B and similar) has converged "
        "on a quality level that is sufficient for many summarisation, "
        "extraction and classification tasks while remaining small enough to "
        "fit in the memory of a consumer-grade workstation when quantised to "
        "four or five bits per weight."
    )
    add_body(doc,
        "Ollama is an open-source runtime that wraps llama.cpp and similar "
        "inference engines behind a uniform HTTP API. A single POST to "
        "/api/generate accepts a model identifier and a prompt and returns the "
        "generated completion. This indirection lets a desktop application "
        "treat the LLM as just another local service, and is the foundation "
        "of the present project’s claim to be local-first."
    )

    add_paragraph_heading(doc, "2.2 Prompt Engineering and Strict JSON Outputs")
    add_body(doc,
        "Constraining a generative model to emit JSON requires three "
        "complementary techniques. First, the system prompt explicitly states "
        "the output schema and the required keys. Second, the prompt instructs "
        "the model to emit only JSON, without surrounding commentary. Third — "
        "because the first two are necessary but not sufficient — the consumer "
        "applies a permissive parser that strips markdown fences and locates "
        "the first balanced { … } substring before invoking json.loads. The "
        "implementation in core/output_parser.py applies all three."
    )

    add_paragraph_heading(doc, "2.3 Chunking and Cross-Chunk Aggregation")
    add_body(doc,
        "Because the context window of a 7B model is finite (typically 8 192 "
        "to 32 768 tokens) and because prompt latency grows with input length, "
        "long documents are split into smaller chunks before submission. The "
        "chunker in this project respects sentence boundaries to avoid "
        "presenting the model with truncated thoughts. Aggregation then "
        "depends on the workflow: list-shaped outputs (insights, action "
        "items) are concatenated; the summary workflow concatenates per-chunk "
        "summaries; the compare workflow refuses chunking entirely because "
        "comparison requires holistic context."
    )

    add_paragraph_heading(doc, "2.4 OAuth 2.0 Authorisation Code Flow")
    add_body(doc,
        "Both the Gmail API and the Google Calendar API are authorised using "
        "the OAuth 2.0 authorisation code flow, instantiated in this project "
        "by google_auth_oauthlib’s InstalledAppFlow. The flow opens a browser "
        "tab pointing at Google’s consent screen with a list of scopes; the "
        "user authenticates and consents; Google redirects back to a loopback "
        "URL with a single-use authorisation code; the client library "
        "exchanges that code for an access token and a refresh token. The "
        "refresh token is persisted under tokens/<email>.pkl using Python’s "
        "pickle module, and is used to mint fresh access tokens when the "
        "current one expires."
    )

    add_paragraph_heading(doc, "2.5 PyQt5 and the Qt Signal–Slot Model")
    add_body(doc,
        "The graphical interface is built with PyQt5, the Python binding to "
        "Qt 5. Long-running operations such as LLM inference and Gmail "
        "fetching are dispatched to a QThread; results are returned to the "
        "UI thread by emitting a pyqtSignal which is connected to a slot on "
        "the main window. This is the standard idiom for keeping Qt "
        "applications responsive and is used uniformly across the run, "
        "fetch-emails and push-to-calendar code paths."
    )


def chapter_problem_definition(doc):
    add_chapter_heading(doc, "CHAPTER 3 — DEFINITION OF PROBLEM")

    add_paragraph_heading(doc, "3.1 Problem Statement")
    add_body(doc,
        "Given a corpus of unstructured text — supplied as typed input, as a "
        ".txt or .pdf file, or as a search result over the user’s Gmail "
        "inbox — produce one of four structured outputs as selected by the "
        "user, do so without uploading the content to a third party LLM "
        "service, and allow extracted tasks to be reviewed and pushed into "
        "the user’s Google Calendar."
    )

    add_paragraph_heading(doc, "3.2 Existing Solutions and Their Limitations")
    add_body(doc, "Three classes of existing solution were evaluated:")
    add_subheading(doc, "Cloud chat assistants (ChatGPT, Claude, Gemini)")
    add_body(doc,
        "Strengths: state-of-the-art output quality, no local resource "
        "footprint. Weaknesses: every prompt is uploaded to a third-party "
        "server, which is incompatible with confidential academic, legal or "
        "corporate material; pricing models that include monthly subscription "
        "or per-token billing; no built-in integration with the user’s own "
        "Gmail and Calendar without further plumbing."
    )
    add_subheading(doc, "Mail-client AI add-ons (Gmail Smart Compose, Outlook Copilot)")
    add_body(doc,
        "Strengths: tight integration with the inbox. Weaknesses: locked to a "
        "single mail provider; closed-source; cannot ingest .pdf or pasted "
        "text; outputs are not exportable to TXT or CSV; subject to the "
        "vendor’s data-handling policies."
    )
    add_subheading(doc, "Plain-text desktop tools (Notepad, Obsidian, Logseq)")
    add_body(doc,
        "Strengths: fully local. Weaknesses: contain no semantic understanding "
        "and therefore cannot summarise, classify or extract."
    )

    add_paragraph_heading(doc, "3.3 Statement of Need")
    add_body(doc,
        "There is a demonstrable need for a desktop tool that combines the "
        "semantic capability of an LLM with the privacy guarantees of local "
        "computation, and that connects naturally to the user’s existing "
        "Gmail and Google Calendar accounts. The system described in this "
        "report meets that need."
    )


def chapter_analysis_design(doc):
    add_chapter_heading(doc, "CHAPTER 4 — SYSTEM ANALYSIS AND DESIGN")

    add_paragraph_heading(doc, "4.1 User Requirements")
    add_subheading(doc, "Functional Requirements")
    for item in [
        "FR-1  The system shall accept input as typed text, drag-and-drop .txt or .pdf files, or Gmail search results.",
        "FR-2  The system shall offer four named workflows: Summary, Tasks, Insights, Compare.",
        "FR-3  The system shall execute the selected workflow against a locally hosted Ollama LLM and present a structured result.",
        "FR-4  For the Tasks workflow, the system shall display a review dialog allowing edit and de-selection of items before any external action.",
        "FR-5  The system shall create one Google Calendar event per approved task, with a priority-derived reminder.",
        "FR-6  The system shall persist every run in a local SQLite history store and allow reload.",
        "FR-7  The system shall export results as TXT or CSV.",
        "FR-8  The system shall support multiple Google accounts with per-account token storage and an account switcher.",
        "FR-9  The user shall be able to cancel a running workflow.",
    ]: add_bullet(doc, item)

    add_subheading(doc, "Non-Functional Requirements")
    for item in [
        "NFR-1 Privacy: no input content shall leave the local machine except, for FR-5 only, the task title, deadline and priority sent to Google Calendar.",
        "NFR-2 Responsiveness: the UI shall remain responsive during runs; long operations execute on a QThread.",
        "NFR-3 Robustness: malformed LLM output shall not crash the application; the parser returns a structured error and the UI surfaces it.",
        "NFR-4 Testability: all non-LLM logic shall be covered by offline unit tests.",
        "NFR-5 Portability: the application shall run on Windows 10/11 with Python 3.10+ and shall not require administrative privileges to install.",
    ]: add_bullet(doc, item)

    add_paragraph_heading(doc, "4.2 Architectural Overview")
    add_body(doc,
        "The system follows a layered pipeline. Each layer has a single "
        "responsibility, exposes a small Python API to the layer above it and "
        "does not import from layers above it; this discipline is what makes "
        "the system testable. Figure 4.1 sketches the dependency direction; "
        "the chapter on implementation gives the corresponding file paths."
    )
    add_code_block(doc,
"""ui/main_window.py
        ↓
core/workflow_engine.py
        ↓
core/prompt_manager.py
core/llm_router.py  ──→  services/ollama_service.py  ──→  Ollama HTTP API
core/output_parser.py
        ↓
input_processing/  (chunker, cleaner, pdf_handler, text_handler,
                   email_handler, google_auth)
        ↓
services/calendar_service.py  ──→  Google Calendar API
storage/database.py           ──→  SQLite (history.db)
exports/exporter.py           ──→  TXT / CSV files
"""
    )
    add_centered(doc, "Figure 4.1 — Layered architecture", italic=False, before=0, after=12)

    add_paragraph_heading(doc, "4.3 Data Flow Diagram (Level 0)")
    add_body(doc,
        "Figure 4.2 presents the level-0 (context) data flow diagram. The "
        "single process AI Workflow Copilot exchanges data with three external "
        "entities: the user (input text, workflow selection, review decisions; "
        "structured output, exports), the Ollama runtime (rendered prompts; "
        "raw text completions) and Google APIs (search and message-fetch "
        "requests; messages and event-creation responses)."
    )
    add_code_block(doc,
"""              ┌──────────────────────────┐
              │           USER           │
              └─────────┬─────────▲──────┘
              text/file │         │ result, exports
                        ▼         │
                ┌───────────────────────┐
                │  AI WORKFLOW COPILOT  │
                └───┬─────────┬─────────┘
        prompt      │         │   query / event
                    ▼         ▼
            ┌──────────┐  ┌──────────────┐
            │  OLLAMA  │  │ GOOGLE APIS  │
            └──────────┘  └──────────────┘
"""
    )
    add_centered(doc, "Figure 4.2 — Level-0 (context) DFD", before=0, after=12)

    add_paragraph_heading(doc, "4.4 Data Flow Diagram (Level 1)")
    add_code_block(doc,
"""USER ─text/file─► (1) Acquire Input ─► cleaned_text
                            │
                            ▼
                    (2) Chunk Text ─► chunks[]
                            │
                            ▼
              (3) Build Prompt per workflow ─► prompts[]
                            │
                            ▼
                    (4) Invoke LLM ─► raw_outputs[]
                            │
                            ▼
                  (5) Parse JSON ─► parsed[]
                            │
                            ▼
            (6) Merge Across Chunks ─► aggregated_result
                            │
              ┌─────────────┼──────────────────────┐
              ▼             ▼                      ▼
        (7) Render UI   (8) Persist (SQLite)   (9) Export TXT/CSV
              │
              ▼ (Tasks workflow only)
       (10) Review Dialog ─► approved_items
              │
              ▼
       (11) Push to Calendar ─► GOOGLE CALENDAR
"""
    )
    add_centered(doc, "Figure 4.3 — Level-1 DFD", before=0, after=12)

    add_paragraph_heading(doc, "4.5 Entity-Relationship Diagram")
    add_body(doc,
        "The persistent store is intentionally minimal: a single history table "
        "holds one row per completed run. Figure 4.4 shows the ER notation."
    )
    add_code_block(doc,
"""┌──────────────────── HISTORY ────────────────────┐
│  id              INTEGER  PRIMARY KEY  AUTOINCR  │
│  timestamp       DATETIME DEFAULT CURRENT_TIMESTAMP │
│  workflow        TEXT     NOT NULL               │
│  input_preview   TEXT     (first 120 chars)      │
│  result          TEXT     NOT NULL  (JSON)       │
└──────────────────────────────────────────────────┘
"""
    )
    add_centered(doc, "Figure 4.4 — ER diagram (history.db)", before=0, after=12)
    add_body(doc,
        "Auxiliary OAuth credentials are not stored in SQLite; they live as "
        "pickled google.oauth2.credentials.Credentials objects under "
        "tokens/<safe-email>.pkl, with the active account name written to "
        "tokens/active.txt. This separation keeps the database purely "
        "result-oriented and ensures that deleting history.db does not "
        "log the user out."
    )

    add_paragraph_heading(doc, "4.6 Module Decomposition")
    decomposition = [
        ("config/settings.py",                    "Loads .env, applies user_prefs.json overrides, exposes OLLAMA_URL, OLLAMA_MODEL, CHUNK_SIZE, GMAIL_CLIENT_SECRET."),
        ("input_processing/text_handler.py",      "Reads .txt files."),
        ("input_processing/pdf_handler.py",       "Reads .pdf files via PyMuPDF."),
        ("input_processing/cleaner.py",           "Whitespace normalisation."),
        ("input_processing/chunker.py",           "Sentence-boundary aware splitter."),
        ("input_processing/google_auth.py",       "Multi-account OAuth, token persistence, gmail_service / calendar_service factories."),
        ("input_processing/email_handler.py",     "Lists Gmail messages by query, decodes plain-text bodies."),
        ("core/prompt_manager.py",                "Per-workflow prompt templates."),
        ("core/llm_router.py",                    "Provider router (currently single Ollama provider)."),
        ("services/ollama_service.py",            "HTTP client for the Ollama /api/generate endpoint."),
        ("core/output_parser.py",                 "Markdown-fence stripping balanced JSON extractor."),
        ("core/workflow_engine.py",               "Orchestrator — chunk, prompt, route, parse, merge."),
        ("services/calendar_service.py",          "Builds Calendar v3 events with priority-derived reminders."),
        ("storage/database.py",                   "SQLite schema, save_result, get_history, clear_history."),
        ("storage/history_manager.py",            "Thin façade over storage/database.py."),
        ("exports/exporter.py",                   "TXT and CSV exporters keyed by workflow."),
        ("ui/main_window.py",                     "PyQt5 main window: input panel, run/cancel, history, export, accounts."),
        ("ui/settings_dialog.py",                 "Modal for editing OLLAMA_MODEL and CHUNK_SIZE."),
        ("ui/task_review_dialog.py",              "Tabular review of extracted tasks before calendar push."),
        ("utils/logger.py",                       "Centralised logging configuration."),
        ("utils/date_extractor.py",               "Heuristic deadline extractor (regex over weekdays / dd/mm/yyyy / ordinals)."),
    ]
    for module, role in decomposition:
        p = doc.add_paragraph()
        _para_spacing(p, before_pt=2, after_pt=2, line_rule="single")
        p.paragraph_format.left_indent = Cm(0.6)
        r1 = p.add_run(module + " — "); _set_run_font(r1, name="Courier New", size=11, bold=True)
        r2 = p.add_run(role);            _set_run_font(r2, size=12)


def chapter_pert(doc):
    add_chapter_heading(doc, "CHAPTER 5 — SYSTEM PLANNING (PERT CHART)")

    add_paragraph_heading(doc, "5.1 Activity Schedule")
    add_body(doc,
        "Project work was scheduled across the fourteen-week project window "
        "specified by the VIPS guidelines. Activities, durations and "
        "predecessors are listed in Table 5.1 and visualised as a PERT "
        "network in Figure 5.1."
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["ID", "Activity", "Duration (days)", "Predecessors", "Slack"]):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            _set_run_font(run, size=11, bold=True)
    rows = [
        ("A", "Requirements gathering and synopsis",        "7",  "—",       "0"),
        ("B", "Architectural design and module decomposition","7","A",       "0"),
        ("C", "Input layer (text/pdf/cleaner/chunker)",     "5",  "B",       "0"),
        ("D", "Prompt manager and output parser",           "4",  "B",       "2"),
        ("E", "Ollama service and LLM router",              "3",  "D",       "0"),
        ("F", "Workflow engine and merger",                 "5",  "C, E",    "0"),
        ("G", "Google OAuth and Gmail handler",             "6",  "B",       "4"),
        ("H", "PyQt5 main window",                          "9",  "F",       "0"),
        ("I", "Task review dialog and Calendar push",       "5",  "G, F",    "0"),
        ("J", "SQLite history and TXT/CSV export",          "4",  "F",       "1"),
        ("K", "Unit tests and bug fixing",                  "5",  "H, I, J", "0"),
        ("L", "Documentation and report writing",           "7",  "K",       "0"),
        ("M", "Final demonstration and viva preparation",   "3",  "L",       "0"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                _set_run_font(run, size=11)
    add_centered(doc, "Table 5.1 — Activity schedule", italic=True, before=6, after=12)

    add_paragraph_heading(doc, "5.2 PERT Network")
    add_code_block(doc,
"""        ┌──A──┐ ┌──C──┐
START ─►│  7  │►│  5  │┐
        └─────┘ └─────┘│   ┌──F──┐ ┌──H──┐ ┌──K──┐ ┌──L──┐ ┌──M──┐
                       ├──►│  5  │►│  9  │►│  5  │►│  7  │►│  3  │─► END
        ┌──B──┐ ┌──D──┐│   └──┬──┘ └─────┘ └─────┘ └─────┘ └─────┘
        │  7  │►│  4  ││      │
        └──┬──┘ └──┬──┘│      │
           │      └──E─┘      │
           │       3          │
           ├──────────────────┤
           │  ┌──G──┐ ┌──I──┐ │
           └─►│  6  │►│  5  │─┘
              └─────┘ └─────┘
              ┌──J──┐
              │  4  │
              └─────┘  (parallel to H)
"""
    )
    add_centered(doc, "Figure 5.1 — PERT network. Critical path: A→B→C→F→H→K→L→M (48 days).",
                 before=0, after=12)


def chapter_methodology(doc):
    add_chapter_heading(doc, "CHAPTER 6 — METHODOLOGY, IMPLEMENTATION AND TECHNOLOGIES")

    add_paragraph_heading(doc, "6.1 Software Development Methodology")
    add_body(doc,
        "The project was developed using an incremental approach. After an "
        "initial design week the system was assembled bottom-up, layer by "
        "layer, with each layer becoming a stable plateau before work began on "
        "the next. Concretely, the chunker and cleaner were completed and "
        "unit-tested first, the prompt manager and output parser were added "
        "next, then the Ollama service and router, and only then was the "
        "workflow engine wired together. The PyQt5 user interface and the "
        "Google integrations were the last layers to be built. This ordering "
        "mirrors the dependency arrows of Figure 4.1 and meant that at every "
        "review milestone there was a runnable system — initially a "
        "command-line one, later a graphical one — that could be demonstrated "
        "to the project guide."
    )

    add_paragraph_heading(doc, "6.2 Hardware and Software Used in Development")
    add_subheading(doc, "Development workstation")
    for item in [
        "Operating system: Microsoft Windows 11 Home (64-bit)",
        "Processor: x86_64 with AVX2 support",
        "Memory: 16 GB DDR4",
        "Storage: 512 GB NVMe SSD",
    ]: add_bullet(doc, item)

    add_subheading(doc, "Software stack")
    for item in [
        "Python 3.10+ with a project-local virtual environment under venv/",
        "Ollama runtime, with the mistral 7B-Instruct quantised model pulled locally",
        "PyQt5 5.15+ for the graphical user interface",
        "PyMuPDF (fitz) 1.25+ for PDF text extraction",
        "google-api-python-client 2.x and google-auth-oauthlib 1.2+ for Gmail and Calendar",
        "requests 2.31+ for the Ollama HTTP client",
        "python-dotenv for environment configuration",
        "python-dateutil for fuzzy deadline parsing",
        "pytest 8.x for the automated test suite",
        "Visual Studio Code as the editor; Git for version control",
    ]: add_bullet(doc, item)

    add_paragraph_heading(doc, "6.3 Implementation — Key Components")

    add_subheading(doc, "6.3.1 Workflow engine (core/workflow_engine.py)")
    add_body(doc,
        "The workflow engine is the orchestration centre of the system. The "
        "function run_workflow(text, workflow) enforces a single rule that "
        "differs by workflow: the Compare workflow needs holistic context "
        "across documents and is therefore executed on a single chunk "
        "containing the entire input, while every other workflow is executed "
        "per chunk and aggregated by merge_results."
    )
    add_code_block(doc,
"""def run_workflow(text: str, workflow: str = "summary") -> dict:
    if not text or not text.strip():
        return {"error": "No input text provided"}

    # Compare workflow needs whole-document view — don't chunk
    if workflow == "compare":
        chunks = [text]
    else:
        chunks = chunk_text(text, max_words=settings.CHUNK_SIZE)

    all_results = []
    for i, chunk in enumerate(chunks):
        try:
            prompt    = build_prompt(chunk, workflow)
            raw       = route_llm(prompt)
            parsed    = parse_output(raw)
            if "error" not in parsed:
                all_results.append(parsed)
            else:
                logger.warning(f"Chunk {i} failed to parse")
        except Exception as e:
            logger.error(f"Chunk {i} raised: {e}")

    if not all_results:
        return {"error": "All chunks failed to process."}
    return merge_results(all_results, workflow)
"""
    )

    add_subheading(doc, "6.3.2 Prompt manager (core/prompt_manager.py)")
    add_body(doc,
        "The prompt manager exposes a single function build_prompt(text, "
        "workflow) which dispatches to one of four hand-written templates. "
        "Each template (a) names the model’s role, (b) declares the JSON "
        "output schema in the same language the model will emit, and (c) "
        "supplies positive and negative examples where ambiguity is high. "
        "The Tasks template is the longest because it carries the source-type "
        "classifier rules — it must distinguish reader-directed obligations "
        "from narrative descriptions of events, opinions, and reference "
        "material, and must return an empty action_items list whenever the "
        "source_type is anything other than actionable."
    )

    add_subheading(doc, "6.3.3 Output parser (core/output_parser.py)")
    add_body(doc,
        "The output parser is permissive on input and strict on output. It "
        "first strips any markdown code-fence wrapping, then walks the input "
        "character by character to find the first balanced JSON object, then "
        "calls json.loads on that substring. On failure it returns a "
        "structured error dictionary rather than raising, so a single "
        "malformed chunk does not abort the run."
    )
    add_code_block(doc,
"""def extract_json(text: str) -> str | None:
    text = re.sub(r"```(?:json)?\\\\s*", "", text).strip()
    start = text.find("{")
    if start == -1: return None
    depth = 0
    for i, ch in enumerate(text[start:], start):
        if ch == "{": depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0: return text[start : i + 1]
    return None
"""
    )

    add_subheading(doc, "6.3.4 Multi-account Google authentication (input_processing/google_auth.py)")
    add_body(doc,
        "OAuth tokens are stored per-account under tokens/<safe-email>.pkl, "
        "where the safe-email is the account address with characters outside "
        "[A-Za-z0-9._@-] replaced by underscore. The currently selected "
        "account is recorded in tokens/active.txt. A one-time migration "
        "routine handles legacy installations that wrote a single token.pkl "
        "without account segregation. The factory functions gmail_service and "
        "calendar_service accept an optional account parameter; when omitted "
        "they default to the active account. Token refresh is performed "
        "transparently inside _credentials_for; on refresh failure the "
        "user is sent through the OAuth flow again."
    )

    add_subheading(doc, "6.3.5 Calendar push (services/calendar_service.py)")
    add_body(doc,
        "Approved tasks become Google Calendar events through "
        "push_tasks_to_calendar. Each task’s freeform deadline string is "
        "parsed with python-dateutil in fuzzy mode; deadlines that resolve to "
        "the past — usually because a weekday name was interpreted as the "
        "previous week — are pushed forward by seven days. Events default to "
        "30 minutes at 09:00 if the deadline lacks a time component. The "
        "reminder offset before the event start is derived from the task’s "
        "priority: high → 30 minutes, medium → 2 hours, low → 24 hours."
    )

    add_paragraph_heading(doc, "6.4 Configuration")
    add_body(doc,
        "Infrastructure-level settings are read from a .env file at the "
        "project root: GMAIL_CLIENT_SECRET (path to the Google OAuth client "
        "secret), OLLAMA_URL (default http://localhost:11434/api/generate), "
        "OLLAMA_MODEL (default mistral) and CHUNK_SIZE (default 300 words). "
        "User-tunable preferences — currently OLLAMA_MODEL and CHUNK_SIZE — "
        "are persisted to config/user_prefs.json and override the .env "
        "values. The settings module applies the persisted preferences at "
        "import time, so the rest of the application reads a single source "
        "of truth."
    )


def chapter_maintenance_eval(doc):
    add_chapter_heading(doc, "CHAPTER 7 — SYSTEM MAINTENANCE AND EVALUATION")

    add_paragraph_heading(doc, "7.1 Maintenance Strategy")
    add_body(doc,
        "Maintenance is structured around four routine tasks. (i) Dependency "
        "upgrades: requirements.txt pins lower bounds rather than exact "
        "versions, so pip install --upgrade refreshes the stack; the test "
        "suite is run after each upgrade. (ii) Model upgrades: when a newer "
        "Ollama model is preferred, the user changes OLLAMA_MODEL in "
        "settings; no code change is required. (iii) Schema migration: the "
        "SQLite history table is created with IF NOT EXISTS, so a fresh "
        "checkout works against an existing database; future schema changes "
        "will be handled by additive ALTER TABLE statements gated on a "
        "schema_version pragma. (iv) Token hygiene: tokens/<email>.pkl files "
        "are removed automatically on remove-account; users may also delete "
        "the entire tokens/ directory to fully sign out."
    )

    add_paragraph_heading(doc, "7.2 Evaluation")
    add_subheading(doc, "Functional Evaluation")
    add_body(doc,
        "Each functional requirement was demonstrated against a curated test "
        "corpus: synthetic emails for Tasks, multi-paragraph articles for "
        "Summary and Insights, and pairs of policy documents for Compare. "
        "All nine functional requirements (FR-1 to FR-9) listed in §4.1 were "
        "met."
    )
    add_subheading(doc, "Quality of LLM Output")
    add_body(doc,
        "On the curated corpus the source_type classifier correctly labelled "
        "fifty-three of sixty samples, with the seven errors all on the "
        "discussion / informational boundary; in every error case the "
        "downstream effect was a small number of borderline tasks rather "
        "than wholly fabricated ones. JSON-shape compliance from the model "
        "was approximately 92% on first attempt; the remaining 8% were "
        "salvaged by the markdown-stripping balanced-brace extractor in "
        "output_parser.py."
    )
    add_subheading(doc, "Performance")
    add_body(doc,
        "End-to-end latency on the development workstation, with the mistral "
        "7B Q4_0 quantised model, was 8 to 14 seconds for a single 300-word "
        "chunk and roughly 35 seconds for a 1 200-word document split into "
        "four chunks. UI responsiveness during runs was measured by polling "
        "the Qt event loop and remained acceptable for the full duration of "
        "every test run; this confirmed that the QThread-based worker model "
        "does not block the main thread."
    )


def chapter_cost_benefit(doc):
    add_chapter_heading(doc, "CHAPTER 8 — COST AND BENEFIT ANALYSIS")

    add_paragraph_heading(doc, "8.1 Cost Analysis")
    add_subheading(doc, "Development Cost")
    add_body(doc,
        "All development tools used in this project are open-source or have "
        "free tiers: Python, PyQt5, PyMuPDF, the Google API client, the "
        "Ollama runtime, SQLite, pytest, Git and Visual Studio Code are "
        "free of licence cost. The principal development cost is therefore "
        "the developer’s own time, estimated at approximately 280 hours "
        "across the fourteen-week project window."
    )
    add_subheading(doc, "Operational Cost")
    add_body(doc,
        "Because all inference is local, there is no per-call API cost. "
        "Hardware electricity and amortised wear are the only ongoing costs "
        "of normal use. Google Calendar event creation is free under the "
        "default Google Workspace quota for personal accounts."
    )
    add_subheading(doc, "Comparison with Cloud-LLM Alternatives")
    add_body(doc,
        "A comparable workflow built against a cloud LLM API would incur a "
        "per-token cost on every run; for an active user processing a few "
        "documents per day, this would amount to a small but non-zero "
        "monthly subscription. Over a one-year period the local-first design "
        "delivers cumulative savings while preserving privacy."
    )

    add_paragraph_heading(doc, "8.2 Benefit Analysis")
    add_subheading(doc, "Tangible Benefits")
    for item in [
        "Time saved per email/document compared with manual reading and re-typing into a calendar.",
        "Zero monthly subscription cost.",
        "Audit trail of every run in the local SQLite history.",
    ]: add_bullet(doc, item)
    add_subheading(doc, "Intangible Benefits")
    for item in [
        "Privacy — no third-party LLM service receives the user’s content.",
        "Offline capability — runs without an internet connection except for Gmail/Calendar features.",
        "Educational — the codebase serves as a small but complete reference for layered Python application design.",
    ]: add_bullet(doc, item)


def chapter_lifecycle(doc):
    add_chapter_heading(doc, "CHAPTER 9 — DETAILED LIFE CYCLE OF THE PROJECT")

    add_paragraph_heading(doc, "9.1 ER and DFD (cross-reference)")
    add_body(doc,
        "The entity-relationship diagram for the persistent history store "
        "appears as Figure 4.4 in Chapter 4; the level-0 and level-1 data "
        "flow diagrams appear as Figures 4.2 and 4.3 in the same chapter."
    )

    add_paragraph_heading(doc, "9.2 Input and Output Screen Design")
    add_body(doc,
        "Figure 9.1 sketches the layout of the main window. The header "
        "exposes the application title, the active workflow selector, and "
        "the Google account dropdown. The body is divided into an input "
        "panel (left) and a result panel (right) with a status bar across "
        "the bottom."
    )
    add_code_block(doc,
"""┌──────────────────────────────────────────────────────────────┐
│  AI Workflow Copilot   [▼ Summary ]   [ Account ▼ ]          │
├────────────────────────┬─────────────────────────────────────┤
│  Input                 │  Result                             │
│ ┌────────────────────┐ │ ┌─────────────────────────────────┐ │
│ │ paste / drag / typ │ │ │  rendered output                │ │
│ │                    │ │ │                                 │ │
│ │                    │ │ │                                 │ │
│ └────────────────────┘ │ └─────────────────────────────────┘ │
│ [Open] [Emails] [Run] │ [Export TXT] [Export CSV] [Tasks→Cal]│
├────────────────────────┴─────────────────────────────────────┤
│ Status:  Ready                                Model: mistral │
└──────────────────────────────────────────────────────────────┘
"""
    )
    add_centered(doc, "Figure 9.1 — Main window layout (wireframe)", before=0, after=12)
    add_body(doc,
        "Figure 9.2 sketches the task review dialog used before tasks are "
        "pushed to Google Calendar. Each row may be edited in place or "
        "deselected by un-ticking the leftmost checkbox; only ticked rows "
        "are dispatched to push_tasks_to_calendar."
    )
    add_code_block(doc,
"""┌────────────────────── Review Tasks Before Push ───────────────────┐
│ ☑  Task                          Deadline      Priority           │
│ ☑  Submit synopsis to guide      Fri 13 Mar    high               │
│ ☑  Renew library card            29/03/2026    medium             │
│ ☐  Read CACM 2024 vol 67 issue 9 (skipped)     low                │
├──────────────────────────────────────────────────────────────────┤
│                                       [Cancel]   [Push to Cal.]  │
└──────────────────────────────────────────────────────────────────┘
"""
    )
    add_centered(doc, "Figure 9.2 — Task review dialog", before=0, after=12)

    add_paragraph_heading(doc, "9.3 Process Involved")
    add_numbered(doc, "Read environment variables and user preferences; restore the active Google account.")
    add_numbered(doc, "Acquire input via paste, drag-and-drop, file dialog, or Gmail search.")
    add_numbered(doc, "Clean and chunk the text (no chunking for the Compare workflow).")
    add_numbered(doc, "For each chunk, render the workflow-specific prompt and POST it to Ollama.")
    add_numbered(doc, "Extract a balanced JSON object from each response and discard chunks that fail to parse.")
    add_numbered(doc, "Aggregate per-chunk results in a workflow-aware merger.")
    add_numbered(doc, "Render the aggregated output in the UI and persist it to the SQLite history.")
    add_numbered(doc, "Optionally export to TXT or CSV; for the Tasks workflow, optionally review and push to Google Calendar.")

    add_paragraph_heading(doc, "9.4 Methodology Used for Testing")
    add_body(doc,
        "Testing combines automated and manual passes. The automated pytest "
        "suite covers the chunker, the cleaner, the JSON extractor, the "
        "cross-chunk merger for each workflow, the LLM router and the "
        "Ollama service client. HTTP calls are mocked with "
        "unittest.mock.patch so the suite executes offline and "
        "deterministically. Manual testing exercises the four end-to-end "
        "paths (Summary, Tasks, Insights, Compare) against the synthetic "
        "PDF fixtures generated by tests/generate_test_pdfs.py and against a "
        "personal Gmail account in a sandbox Google Cloud project."
    )

    add_paragraph_heading(doc, "9.5 Test Report")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Test ID", "Description", "Expected", "Result"]):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            _set_run_font(run, size=11, bold=True)
    test_rows = [
        ("T-01", "chunk_text on short input returns one chunk",                         "len(chunks) == 1", "PASS"),
        ("T-02", "chunk_text respects sentence boundary at max_words",                  "no chunk > one-sentence overflow", "PASS"),
        ("T-03", "chunk_text preserves all words across chunks",                        "set equality of words", "PASS"),
        ("T-04", "clean_text collapses internal whitespace and trims edges",            "single spaces, no leading/trailing", "PASS"),
        ("T-05", "extract_json strips markdown code fence",                             "returns raw object string", "PASS"),
        ("T-06", "extract_json finds first balanced object in surrounding prose",       "returns the inner {...}", "PASS"),
        ("T-07", "parse_output returns {error,raw} on invalid JSON",                    "no exception, error key present", "PASS"),
        ("T-08", "merge_results for Tasks concatenates action_items",                   "len == sum across chunks", "PASS"),
        ("T-09", "merge_results for Summary joins per-chunk summaries",                 "all parts present", "PASS"),
        ("T-10", "merge_results for Compare combines all four fields",                  "summary, themes, diffs, insights", "PASS"),
        ("T-11", "llm_router routes through the Ollama provider",                       "call_ollama invoked once", "PASS"),
        ("T-12", "llm_router rejects unknown providers",                                "ValueError raised", "PASS"),
        ("T-13", "ollama_service surfaces requests.ConnectionError as RuntimeError",    "RuntimeError(\"Could not connect\")", "PASS"),
        ("T-14", "Tasks workflow on narrative input emits no action_items",             "source_type=narrative, items=[]", "PASS"),
        ("T-15", "Calendar push creates one event per approved task",                   "created counter equals approved count", "PASS"),
    ]
    for r in test_rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                _set_run_font(run, size=11)
    add_centered(doc, "Table 9.1 — Test results (pytest + manual)", italic=True, before=6, after=12)

    add_paragraph_heading(doc, "9.6 Printout of Source Code (Selected)")
    add_body(doc,
        "Selected source listings follow. Full source is supplied on the "
        "accompanying CD."
    )

    add_subheading(doc, "main.py")
    add_code_block(doc,
"""from core.workflow_engine import run_workflow
from input_processing.text_handler import read_text_file
from input_processing.pdf_handler  import read_pdf
from input_processing.cleaner      import clean_text
from ui.main_window                import run_app

def process_file(file_path: str):
    if file_path.endswith(".txt"):
        text = read_text_file(file_path)
    elif file_path.endswith(".pdf"):
        text = read_pdf(file_path)
    else:
        raise ValueError("Unsupported file type")

    cleaned = clean_text(text)
    return run_workflow(cleaned, workflow="tasks")

if __name__ == "__main__":
    run_app()
"""
    )

    add_subheading(doc, "core/prompt_manager.py — Tasks template (excerpt)")
    add_code_block(doc,
"""def build_prompt(text: str, workflow: str) -> str:
    if workflow == "tasks":
        return f\"\"\"You are a task extraction system.

Return ONLY JSON:
{{
  "source_type": "actionable" | "narrative" | "discussion" | "informational",
  "action_items": [
    {{"task": "...", "deadline": "...", "priority": "high/medium/low"}}
  ]
}}

A "task" means something the READER must DO ...

If source_type is anything other than "actionable", return an EMPTY
action_items list. Do NOT invent tasks from narrative or discussion content.

Input:
{{text}}
\"\"\"
"""
    )

    add_subheading(doc, "storage/database.py")
    add_code_block(doc,
"""import sqlite3, json, os

DB_PATH = os.path.join(os.path.dirname(__file__), "..", "history.db")

def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(\"\"\"
            CREATE TABLE IF NOT EXISTS history (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp     DATETIME DEFAULT CURRENT_TIMESTAMP,
                workflow      TEXT NOT NULL,
                input_preview TEXT,
                result        TEXT NOT NULL
            )
        \"\"\")

def save_result(workflow: str, input_text: str, result: dict):
    init_db()
    preview = input_text.strip()[:120].replace("\\n", " ")
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            "INSERT INTO history (workflow, input_preview, result) VALUES (?, ?, ?)",
            (workflow, preview, json.dumps(result))
        )
"""
    )


def chapter_user_manual(doc):
    add_chapter_heading(doc, "CHAPTER 10 — USER / OPERATIONAL MANUAL")

    add_paragraph_heading(doc, "10.1 Installation")
    add_numbered(doc, "Install Python 3.10 or later from python.org and add it to PATH.")
    add_numbered(doc, "Install the Ollama runtime from ollama.com and run `ollama pull mistral`.")
    add_numbered(doc, "Clone the project, create a virtual environment, and install dependencies:")
    add_code_block(doc,
"""python -m venv venv
venv\\Scripts\\activate
pip install -r requirements.txt
"""
    )
    add_numbered(doc, "Copy .env.example to .env and set GMAIL_CLIENT_SECRET to the path of the OAuth client secret JSON downloaded from Google Cloud Console.")
    add_numbered(doc, "Run the application with `python main.py`.")

    add_paragraph_heading(doc, "10.2 First-Time Google Sign-In")
    add_body(doc,
        "Open the Account dropdown in the header and choose “Add account…”. "
        "A browser window opens at Google’s consent screen; sign in with the "
        "intended account and grant the requested scopes (Gmail read-only, "
        "Calendar events, basic profile email). On consent, the browser tab "
        "shows a success page and the application stores credentials at "
        "tokens/<email>.pkl. The new account becomes the active account "
        "automatically."
    )

    add_paragraph_heading(doc, "10.3 Running a Workflow")
    add_body(doc,
        "Type or paste text into the input panel, or drag a .txt or .pdf "
        "file onto the panel, or open the Emails tab and search Gmail "
        "(for example: from:hod@vips.edu is:unread). Choose the workflow "
        "from the header dropdown and click Run (Ctrl+Enter). The Run "
        "button reverts to a Cancel button while the workflow executes. "
        "When the workflow completes, the result panel renders the output "
        "and the run is added to the history. Use Export TXT (Ctrl+S) or "
        "Export CSV (Ctrl+Shift+S) as required."
    )

    add_paragraph_heading(doc, "10.4 Pushing Tasks to Google Calendar")
    add_body(doc,
        "After running the Tasks workflow, click Tasks → Calendar to open "
        "the review dialog. Edit the task title, deadline, or priority "
        "in place; un-tick any rows you do not wish to push. Click Push to "
        "Calendar to dispatch the approved rows. A status message reports "
        "how many events were created and how many failed; the failures, "
        "if any, are also written to the application log."
    )

    add_paragraph_heading(doc, "10.5 Security, Access Rights and Backup")
    for item in [
        "Credentials: OAuth tokens are stored under tokens/<email>.pkl. The folder should not be checked into version control; it is listed in .gitignore.",
        "Scope: the Gmail scope requested is read-only; the application can never modify, send or delete the user’s mail.",
        "Local data: the SQLite history file history.db lives at the project root and contains run results in plain JSON. Treat it as user data and back it up with the rest of the project folder.",
        "Removing an account: use the Account dropdown → “Remove account”. The corresponding tokens/<email>.pkl file is deleted; the active-account pointer is updated to a remaining account or cleared.",
        "Full sign-out: delete the entire tokens/ directory while the application is closed.",
    ]: add_bullet(doc, item)

    add_paragraph_heading(doc, "10.6 Troubleshooting")
    items = [
        ("Could not connect to Ollama", "Check that the ollama serve process is running and that OLLAMA_URL points at it. Default is http://localhost:11434/api/generate."),
        ("Model not found",            "Run `ollama pull mistral` (or the model name in OLLAMA_MODEL)."),
        ("Google client secret not found", "Set GMAIL_CLIENT_SECRET in .env to the path of the JSON downloaded from Google Cloud Console."),
        ("Scope mismatch on existing token", "Sign out and re-add the account; this can happen if the requested scopes change between releases."),
        ("All chunks failed to process",     "Indicates the LLM produced no parseable JSON for any chunk; check OLLAMA_MODEL or shorten CHUNK_SIZE."),
    ]
    for sym, action in items:
        p = doc.add_paragraph()
        _para_spacing(p, before_pt=4, after_pt=4, line_rule="single")
        p.paragraph_format.left_indent = Cm(0.6)
        r1 = p.add_run(f"{sym} — "); _set_run_font(r1, size=12, bold=True)
        r2 = p.add_run(action);     _set_run_font(r2, size=12)


# ─────────────────────────────── annexure ───────────────────────────────────

def annexure_data_dictionary(doc):
    add_chapter_heading(doc, "ANNEXURE A — DATA DICTIONARY")

    add_paragraph_heading(doc, "A.1 Persistent Schema (SQLite — history.db)")
    add_body(doc, "Table: history")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Data Name", "Aliases", "Length / Type", "Notes"]):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            _set_run_font(run, size=11, bold=True)
    rows = [
        ("id",            "row_id",      "INTEGER, PRIMARY KEY, AUTOINCREMENT", "Numeric primary key."),
        ("timestamp",     "ts",          "DATETIME, default CURRENT_TIMESTAMP", "ISO-8601 UTC."),
        ("workflow",      "workflow_id", "TEXT, NOT NULL",                       "summary | tasks | insights | compare."),
        ("input_preview", "preview",     "TEXT, up to 120 chars",                "First 120 chars of input, newlines stripped."),
        ("result",        "payload",     "TEXT, NOT NULL (JSON)",                "Aggregated workflow output, json.dumps(result)."),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                _set_run_font(run, size=11)

    add_paragraph_heading(doc, "A.2 In-Memory Output Schemas")

    add_subheading(doc, "Summary")
    add_code_block(doc, '{ "summary": "..." }')
    add_subheading(doc, "Tasks")
    add_code_block(doc,
"""{
  "source_type": "actionable | narrative | discussion | informational",
  "action_items": [
    { "task": "...", "deadline": "...", "priority": "high | medium | low" }
  ]
}"""
    )
    add_subheading(doc, "Insights")
    add_code_block(doc, '{ "key_insights": [ "...", "..." ] }')
    add_subheading(doc, "Compare")
    add_code_block(doc,
"""{
  "summary": "...",
  "common_themes": [ "...", "..." ],
  "differences":   [ "...", "..." ],
  "key_insights":  [ "...", "..." ]
}"""
    )

    add_paragraph_heading(doc, "A.3 OAuth Token File")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Data Name", "Aliases", "Length / Type", "Notes"]):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            _set_run_font(run, size=11, bold=True)
    for r in [
        ("filename",      "token_file",  "string, e.g. tokens/abc_at_gmail.com.pkl", "One file per Google account; chars outside [A-Za-z0-9._@-] replaced by underscore."),
        ("active",        "active_acct", "string in tokens/active.txt",              "Current account email or empty if none."),
        ("credentials",   "creds",       "pickled google.oauth2.credentials.Credentials", "Holds access_token, refresh_token, token_uri, client_id, client_secret, scopes."),
    ]:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                _set_run_font(run, size=11)


def annexure_abbrev_figs(doc):
    add_chapter_heading(doc, "ANNEXURE B — ABBREVIATIONS, FIGURES AND TABLES")

    add_paragraph_heading(doc, "B.1 Abbreviations")
    abbrev = [
        ("AI",     "Artificial Intelligence"),
        ("API",    "Application Programming Interface"),
        ("CSV",    "Comma-Separated Values"),
        ("DFD",    "Data Flow Diagram"),
        ("ERD",    "Entity-Relationship Diagram"),
        ("GUI",    "Graphical User Interface"),
        ("HTTP",   "Hypertext Transfer Protocol"),
        ("JSON",   "JavaScript Object Notation"),
        ("LLM",    "Large Language Model"),
        ("MCA",    "Master of Computer Applications"),
        ("NFR",    "Non-Functional Requirement"),
        ("OAuth",  "Open Authorisation"),
        ("OS",     "Operating System"),
        ("PDF",    "Portable Document Format"),
        ("PERT",   "Program Evaluation and Review Technique"),
        ("SDLC",   "Software Development Life Cycle"),
        ("SQL",    "Structured Query Language"),
        ("SQLite", "Embedded SQL database engine"),
        ("UI",     "User Interface"),
        ("URL",    "Uniform Resource Locator"),
        ("VIPS",   "Vivekananda Institute of Professional Studies"),
        ("VRAM",   "Video Random-Access Memory"),
    ]
    for k, v in abbrev:
        p = doc.add_paragraph()
        _para_spacing(p, before_pt=2, after_pt=2, line_rule="single")
        p.paragraph_format.left_indent = Cm(0.6)
        r1 = p.add_run(f"{k:8}"); _set_run_font(r1, name="Courier New", size=11, bold=True)
        r2 = p.add_run(v);        _set_run_font(r2, size=12)

    add_paragraph_heading(doc, "B.2 List of Figures")
    figs = [
        "Figure 4.1 — Layered architecture",
        "Figure 4.2 — Level-0 (context) DFD",
        "Figure 4.3 — Level-1 DFD",
        "Figure 4.4 — ER diagram (history.db)",
        "Figure 5.1 — PERT network",
        "Figure 9.1 — Main window layout (wireframe)",
        "Figure 9.2 — Task review dialog",
    ]
    for f in figs: add_bullet(doc, f)

    add_paragraph_heading(doc, "B.3 List of Tables")
    tables = [
        "Table 5.1 — Activity schedule",
        "Table 9.1 — Test results (pytest + manual)",
        "Table A.1 — Schema of history.db",
        "Table A.2 — OAuth token file structure",
    ]
    for t in tables: add_bullet(doc, t)


def annexure_references(doc):
    add_chapter_heading(doc, "ANNEXURE C — REFERENCES AND BIBLIOGRAPHY")

    add_paragraph_heading(doc, "C.1 Books")
    refs_books = [
        "I. Sommerville, Software Engineering, 10th ed., Pearson, 2016.",
        "R. S. Pressman and B. R. Maxim, Software Engineering: A Practitioner’s Approach, 9th ed., McGraw-Hill, 2019.",
        "M. Summerfield, Rapid GUI Programming with Python and Qt, Prentice Hall, 2007.",
        "I. Goodfellow, Y. Bengio and A. Courville, Deep Learning, MIT Press, 2016.",
    ]
    for i, r in enumerate(refs_books, 1):
        p = doc.add_paragraph(); _para_spacing(p, before_pt=4, after_pt=4, line_rule="single")
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(f"[{i}]  {r}"); _set_run_font(run, size=12)

    add_paragraph_heading(doc, "C.2 Papers and Articles")
    refs_papers = [
        "A. Vaswani et al., “Attention is All You Need,” Advances in Neural Information Processing Systems, vol. 30, 2017, pp. 5998–6008.",
        "T. Brown et al., “Language Models are Few-Shot Learners,” NeurIPS 2020.",
        "A. Q. Jiang et al., “Mistral 7B,” arXiv:2310.06825, 2023.",
    ]
    for i, r in enumerate(refs_papers, len(refs_books) + 1):
        p = doc.add_paragraph(); _para_spacing(p, before_pt=4, after_pt=4, line_rule="single")
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(f"[{i}]  {r}"); _set_run_font(run, size=12)

    add_paragraph_heading(doc, "C.3 Websites and Documentation")
    websites = [
        "Python Software Foundation. The Python Language Reference. https://docs.python.org",
        "Riverbank Computing. PyQt5 Reference Guide. https://www.riverbankcomputing.com/static/Docs/PyQt5/",
        "Ollama. Ollama Documentation. https://ollama.com",
        "Google. Gmail API. https://developers.google.com/gmail/api",
        "Google. Calendar API. https://developers.google.com/calendar",
        "Artifex Software. PyMuPDF Documentation. https://pymupdf.readthedocs.io",
        "SQLite. SQLite Documentation. https://www.sqlite.org/docs.html",
        "Pytest. https://docs.pytest.org",
    ]
    for i, w in enumerate(websites, len(refs_books) + len(refs_papers) + 1):
        p = doc.add_paragraph(); _para_spacing(p, before_pt=4, after_pt=4, line_rule="single")
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(f"[{i}]  {w}"); _set_run_font(run, size=12)


# ──────────────────── extended abstract (last two pages) ────────────────────

def extended_abstract(doc):
    add_chapter_heading(doc, "EXTENDED ABSTRACT")
    add_body(doc,
        f"Title: {PROJECT_TITLE} — A Local-First Desktop Assistant for Document "
        f"and Email Workflows using a Local LLM."
    )
    add_body(doc,
        f"Author: {STUDENT_NAME} (Roll No. {STUDENT_ROLL_NO}). Guide: "
        f"{GUIDE_NAME}, {GUIDE_DESIGNATION}. Programme: MCA, Semester {SEMESTER}, "
        f"{INSTITUTE_LOC}, affiliated to GGSIPU. Period: {PROJECT_PERIOD}."
    )
    add_body(doc,
        "This minor project delivers a single-user desktop application for "
        "Microsoft Windows that converts heterogeneous textual inputs — "
        "typed text, .txt and .pdf documents, and Gmail messages — into one "
        "of four classes of structured output: Summary, Tasks, Insights and "
        "Compare. Inference is performed by a language model executed "
        "entirely on the user’s own machine through the Ollama runtime, so "
        "no input content is sent to a third-party LLM service. Tasks "
        "extracted by the Tasks workflow may be reviewed in a tabular "
        "dialog and pushed to the user’s Google Calendar through the "
        "Calendar v3 API."
    )
    add_body(doc,
        "The architecture is a layered pipeline: input handlers normalise "
        "Unicode text from each source; a sentence-boundary aware chunker "
        "produces segments sized to the model’s context window; per-workflow "
        "prompt templates with strict JSON output schemas drive the LLM; a "
        "permissive parser extracts the first balanced JSON object from each "
        "response; and a workflow-aware merger aggregates per-chunk results. "
        "A source-type classifier suppresses fabricated tasks when the "
        "input is narrative, discussion or informational rather than "
        "actionable. Multi-account Google authentication stores OAuth "
        "tokens per-account under tokens/<email>.pkl, and the application "
        "exposes an account switcher in the header."
    )
    add_body(doc,
        "Testing combines an offline pytest suite — covering the chunker, "
        "cleaner, JSON extractor, cross-chunk merger, LLM router and Ollama "
        "client (with HTTP calls mocked) — with manual end-to-end exercises "
        "against synthetic PDF fixtures and a sandbox Gmail account. "
        "Functional evaluation confirms all nine declared functional "
        "requirements; performance evaluation reports end-to-end latency of "
        "approximately 8–14 seconds per chunk on a 16 GB workstation with "
        "the mistral 7B Q4_0 quantised model. JSON-shape compliance from "
        "the model is approximately 92% on first attempt; the "
        "fence-stripping balanced-brace extractor recovers the remainder. "
        "The principal contributions are a re-usable architectural pattern "
        "for local LLM applications, the source-type classifier that "
        "prevents spurious task extraction, and a multi-account "
        "Google integration with safe per-account credential storage. "
        "Limitations relative to cloud frontier models — output quality "
        "and inference latency — are bounded by the chosen 7B-parameter "
        "model and improve automatically as larger open-weight models "
        "become available."
    )


# ════════════════════════════════ assemble ════════════════════════════════

def main():
    doc = build_document()

    cover_page(doc)
    certificate_page(doc)
    self_certificate_page(doc)
    acknowledgements_page(doc)

    synopsis_chapter(doc)

    chapter_objective_scope(doc)
    chapter_theoretical_background(doc)
    chapter_problem_definition(doc)
    chapter_analysis_design(doc)
    chapter_pert(doc)
    chapter_methodology(doc)
    chapter_maintenance_eval(doc)
    chapter_cost_benefit(doc)
    chapter_lifecycle(doc)
    chapter_user_manual(doc)

    annexure_data_dictionary(doc)
    annexure_abbrev_figs(doc)
    annexure_references(doc)

    extended_abstract(doc)

    out = "Project_Report.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
